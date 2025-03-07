import os
import pdfplumber
from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text

def extract_text(file_path):
    """
    Extract text from a PDF or DOCX file.

    :param file_path: Path to the file
    :return: Extracted text as a string
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError("File not found: " + file_path)

    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")


def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n" if page.extract_text() else ""
    return text.strip()


def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_text_from_file(file):
    """Extract text from PDF or DOCX file."""
    filename = file.filename
    temp_path = os.path.join('uploads', filename)
    os.makedirs('uploads', exist_ok=True)
    
    try:
        # Save the uploaded file temporarily
        file.save(temp_path)
        
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = extract_pdf_text(temp_path)
        elif filename.lower().endswith('.docx'):
            doc = Document(temp_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError('Unsupported file format')
        
        return text.strip()
    finally:
        # Clean up the temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)